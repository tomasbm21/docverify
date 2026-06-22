"""
Test script for intentionally malformed documents.
Tests pipeline robustness against edge cases.
"""
import os
import sys
import traceback
import warnings
from pathlib import Path

# Ensure we can import docverify
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from openpyxl import Workbook


def create_test_files():
    """Create intentionally malformed test documents."""
    test_dir = Path(__file__).parent.parent / "data" / "test_malformed"
    test_dir.mkdir(exist_ok=True)

    # 1. Empty .docx (no paragraphs, no tables)
    empty_docx = Document()
    # Remove all default content
    for p in empty_docx.paragraphs:
        p._element.getparent().remove(p._element)
    empty_docx.save(test_dir / "empty.docx")
    print("Created: empty.docx")

    # 2. .docx with header only (no data)
    header_only = Document()
    header_only.add_heading("BILL OF LADING", level=1)
    header_only.save(test_dir / "header_only.docx")
    print("Created: header_only.docx")

    # 3. .xlsx with empty sheet
    wb_empty = Workbook()
    ws_empty = wb_empty.active
    ws_empty.title = "EmptySheet"
    wb_empty.save(test_dir / "empty_sheet.xlsx")
    print("Created: empty_sheet.xlsx")

    # 4. .xlsx with circular formula references
    wb_circular = Workbook()
    ws_circular = wb_circular.active
    ws_circular.title = "Circular"
    # A1 references B1, B1 references A1
    ws_circular["A1"] = "=B1"
    ws_circular["B1"] = "=A1"
    # Add some more circular chains
    ws_circular["C1"] = "=D1+E1"
    ws_circular["D1"] = "=C1"
    ws_circular["E1"] = "=C1"
    wb_circular.save(test_dir / "circular_refs.xlsx")
    print("Created: circular_refs.xlsx")

    # 5. .docx with mixed encodings (simulate by adding text with different encodings)
    mixed_doc = Document()
    # Add UTF-8 text
    mixed_doc.add_paragraph("UTF-8: Héllo Wörld Ñoño café résumé")
    # Add text that would be Latin-1 specific characters
    mixed_doc.add_paragraph("Latin-1 style: ñ á é í ó ú ü ß æ ø å")
    # Add some problematic characters (Unicode private use area, surrogates)
    mixed_doc.add_paragraph("Special: ‘’“”—…")
    mixed_doc.save(test_dir / "mixed_encoding.docx")
    print("Created: mixed_encoding.docx")

    # 5b. Truly corrupted .docx (valid zip but invalid XML inside)
    import zipfile
    import io
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Write minimal [Content_Types].xml
        zf.writestr('[Content_Types].xml', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>')
        # Write _rels/.rels
        zf.writestr('_rels/.rels', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>')
        # Write word/_rels/document.xml.rels
        zf.writestr('word/_rels/document.xml.rels', '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
        # Write corrupted word/document.xml (truncated + invalid XML)
        zf.writestr('word/document.xml', '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>BILL OF LADING</w:t></w:r></w:p><w:p><w:r><w:t>Truncated content here and then we just cut')
    with open(test_dir / "corrupted.docx", "wb") as f:
        f.write(zip_buf.getvalue())
    print("Created: corrupted.docx (valid zip, truncated XML)")

    # 6. .xlsx with extremely wide sheet (1000 columns)
    wb_wide = Workbook()
    ws_wide = wb_wide.active
    ws_wide.title = "Wide"
    for col in range(1, 1001):
        ws_wide.cell(1, col, f"Col{col}")
    ws_wide.cell(2, 1, "Some value")
    wb_wide.save(test_dir / "wide_sheet.xlsx")
    print("Created: wide_sheet.xlsx")

    # 7. .xlsx with formulas that produce errors
    wb_errors = Workbook()
    ws_errors = wb_errors.active
    ws_errors.title = "Errors"
    ws_errors["A1"] = "Invoice Number"
    ws_errors["B1"] = "=1/0"  # #DIV/0!
    ws_errors["A2"] = "Date"
    ws_errors["B2"] = "=INVALID_FUNCTION()"  # #NAME?
    ws_errors["A3"] = "Amount"
    ws_errors["B3"] = "=VLOOKUP(A1,Sheet2!A:B,2,FALSE)"  # #REF if Sheet2 doesn't exist
    wb_errors.save(test_dir / "formula_errors.xlsx")
    print("Created: formula_errors.xlsx")

    # 8. Zero-byte .docx file (empty file with .docx extension)
    with open(test_dir / "zero_byte.docx", "wb") as f:
        pass
    print("Created: zero_byte.docx (0 bytes)")

    # 9. .docx file that is actually a .txt renamed to .docx
    with open(test_dir / "fake_docx.docx", "w", encoding="utf-8") as f:
        f.write("This is actually a plain text file pretending to be a docx.\n")
        f.write("BILL OF LADING\nInvoice: 12345\nDate: 2025-01-15\n")
    print("Created: fake_docx.docx (plain text with .docx extension)")

    # 10. .xlsx that is actually a CSV renamed
    with open(test_dir / "fake_xlsx.xlsx", "w", encoding="utf-8") as f:
        f.write("Invoice,Date,Amount\n")
        f.write("INV-001,2025-01-15,1000.00\n")
    print("Created: fake_xlsx.xlsx (CSV with .xlsx extension)")

    print(f"\nAll test files created in: {test_dir}")
    return test_dir


def run_pipeline(test_dir):
    """Run the DocVerify pipeline against the malformed test corpus."""
    from docverify.pipeline import run_pipeline as pipeline_run

    out_dir = test_dir / "output"
    out_dir.mkdir(exist_ok=True)

    print(f"\n{'='*60}")
    print(f"Running pipeline against malformed corpus: {test_dir}")
    print(f"Output directory: {out_dir}")
    print(f"{'='*60}\n")

    try:
        result = pipeline_run(
            corpus_dir=str(test_dir),
            out_dir=str(out_dir),
            use_llm=False,
            numeric_tolerance=0.0,
        )
        print(f"\nPipeline completed. Result type: {type(result)}")
        if result:
            print(f"Result: {result}")
        return result, None
    except Exception as e:
        print(f"\n*** PIPELINE CRASHED ***")
        print(f"Exception type: {type(e).__name__}")
        print(f"Exception message: {e}")
        print(f"\nFull traceback:")
        traceback.print_exc()
        return None, e


def test_individual_components(test_dir):
    """Test individual components to isolate failure points."""
    from docverify.ingestion.ingest import ingest_dir, ingest_file

    print(f"\n{'='*60}")
    print("Testing individual components...")
    print(f"{'='*60}\n")

    test_files = list(test_dir.glob("*"))
    test_files = [f for f in test_files if f.is_file()]

    for test_file in test_files:
        print(f"\n--- Testing: {test_file.name} ---")

        # Test ingestion
        try:
            doc = ingest_file(str(test_file))
            print(f"  Ingestion: OK")
            print(f"    doc_id: {doc.doc_id[:16]}...")
            print(f"    text length: {len(doc.text)} chars")
            print(f"    tables: {len(doc.tables)}")
            print(f"    text preview: {repr(doc.text[:100]) if doc.text else '(empty)'}")
        except Exception as e:
            print(f"  Ingestion: FAILED - {type(e).__name__}: {e}")
            traceback.print_exc()


def main():
    """Main test runner."""
    print("="*60)
    print("MALFORMED DOCUMENT TEST SUITE")
    print("="*60)

    # Create test files
    test_dir = create_test_files()

    # Test individual components first
    test_individual_components(test_dir)

    # Run full pipeline
    result, error = run_pipeline(test_dir)

    # Summary
    print(f"\n{'='*60}")
    print("TEST SUMMARY")
    print(f"{'='*60}")
    if error:
        print(f"Pipeline CRASHED with: {type(error).__name__}: {error}")
        print("Result: FAIL")
        return 1
    else:
        print("Pipeline completed without crashing")
        print("Result: PASS (needs manual review of output)")
        return 0


if __name__ == "__main__":
    sys.exit(main())
