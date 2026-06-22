"""Integration test — tiny end-to-end corpus with mixed formats.

OWNER: Agent 92 (QA / Test Engineer)
Tests the full pipeline (Ingestion -> Extraction -> Matching -> Verification ->
Reporting -> Scoring) on a hand-built corpus of 1 clean + 1 broken shipment
using mixed .docx and .xlsx formats.

Fast, deterministic, offline (no LLM calls).
"""

import json
import os
from pathlib import Path

import pytest

from docverify.schemas.models import CanonicalDoc, ShipmentVerdict


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def _make_docx(path: str, title: str, metadata: dict[str, str],
               body: str = "") -> None:
    """Create a minimal .docx with a title, metadata table, and optional body."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)

    table = doc.add_table(rows=len(metadata), cols=2)
    for i, (k, v) in enumerate(metadata.items()):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    if body:
        doc.add_paragraph(body)

    doc.save(path)


def _make_xlsx_invoice(path: str, order_no: str, bl_no: str,
                       container_no: str, net_kg: float, gross_kg: float,
                       cartons: int, value: float, currency: str,
                       shipper: str = "Test Foods S.r.l.") -> None:
    """Create a minimal xlsx invoice with metadata + line items."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Invoice"

    # Metadata block
    ws.append(["Order No.", order_no])
    ws.append(["B/L No.", bl_no])
    ws.append(["Container No.", container_no])
    ws.append(["Shipper", shipper])
    ws.append(["Currency", currency])
    ws.append([])  # blank row

    # Line-item header
    ws.append(["Description", "Lot", "Cartons", "Net Kg", "Gross Kg", "Unit Price", "Amount"])
    # Line items
    ws.append(["Spaghetti No.5", "L001", cartons // 2, net_kg / 2, gross_kg / 2, 5.50, value / 2])
    ws.append(["Penne Rigate", "L002", cartons // 2, net_kg / 2, gross_kg / 2, 6.00, value / 2])
    # Total row
    ws.append(["Total", "", cartons, net_kg, gross_kg, "", value])

    wb.save(path)


def _make_answer_key(tmp_path: Path, shipments: list[dict]) -> Path:
    """Create a minimal answer_key.json for the given shipments."""
    manifest = []
    answer_key = []

    for s in shipments:
        for doc_info in s["documents"]:
            manifest.append({
                "file": doc_info["file"],
                "shipment": s["shipment_num"],
                "type": doc_info["type"],
                "format": doc_info["format"],
            })
        answer_key.append({
            "shipment": s["shipment_num"],
            "order_no": s["order_no"],
            "bl_no": s["bl_no"],
            "container": s["container"],
            "has_planted_discrepancy": s["has_planted_discrepancy"],
            "discrepancy_in_doc_type": s.get("discrepancy_in_doc_type"),
            "discrepancy_detail": s.get("discrepancy_detail", []),
            "ground_truth_totals": s.get("ground_truth_totals", {}),
        })

    ak = {
        "manifest": manifest,
        "answer_key": answer_key,
        "total_documents": len(manifest),
    }

    ak_path = tmp_path / "answer_key.json"
    with open(ak_path, "w") as f:
        json.dump(ak, f)
    return ak_path


# ---------------------------------------------------------------------------
# Integration corpus fixture
# ---------------------------------------------------------------------------


@pytest.fixture()
def mixed_corpus(tmp_path: Path) -> tuple[Path, Path]:
    """Create a tiny corpus with mixed .docx and .xlsx formats.

    Clean shipment (S01):
      - BL (docx): ORD-INT-001, BL-INT-001, CONT-INT-001
      - Invoice (xlsx): ORD-INT-001, BL-INT-001, CONT-INT-001

    Broken shipment (S02):
      - BL (docx): ORD-INT-002, BL-INT-002, CONT-INT-002
      - Invoice (xlsx): ORD-INT-002-ALTERED (planted discrepancy in order_no)
    """
    corpus = tmp_path / "corpus"
    corpus.mkdir()

    # --- Clean shipment (S01) ---
    _make_docx(
        str(corpus / "S01_BL_v1.docx"),
        "BILL OF LADING",
        {
            "Order No.": "ORD-INT-001",
            "B/L No.": "BL-INT-001",
            "Container No.": "CONT-INT-001",
            "Shipper": "Alpha Foods S.r.l.",
            "Consignee": "Beta Imports Ltd.",
        },
    )

    _make_xlsx_invoice(
        str(corpus / "S01_Invoice_v1.xlsx"),
        order_no="ORD-INT-001",
        bl_no="BL-INT-001",
        container_no="CONT-INT-001",
        net_kg=4200.0,
        gross_kg=4550.0,
        cartons=350,
        value=15000.00,
        currency="EUR",
        shipper="Alpha Foods S.r.l.",
    )

    # --- Broken shipment (S02) ---
    _make_docx(
        str(corpus / "S02_BL_v1.docx"),
        "BILL OF LADING",
        {
            "Order No.": "ORD-INT-002",
            "B/L No.": "BL-INT-002",
            "Container No.": "CONT-INT-002",
            "Shipper": "Gamma Pasta Co.",
            "Consignee": "Delta Trading FZE",
        },
    )

    _make_xlsx_invoice(
        str(corpus / "S02_Invoice_v1.xlsx"),
        order_no="ORD-INT-002-ALTERED",  # planted discrepancy
        bl_no="BL-INT-002",
        container_no="CONT-INT-002",
        net_kg=3100.0,
        gross_kg=3400.0,
        cartons=250,
        value=12000.00,
        currency="USD",
        shipper="Gamma Pasta Co.",
    )

    # Answer key
    ak_path = _make_answer_key(tmp_path, [
        {
            "shipment_num": 1,
            "order_no": "ORD-INT-001",
            "bl_no": "BL-INT-001",
            "container": "CONT-INT-001",
            "has_planted_discrepancy": False,
            "documents": [
                {"file": "S01_BL_v1.docx", "type": "bill_of_lading", "format": "docx"},
                {"file": "S01_Invoice_v1.xlsx", "type": "invoice", "format": "xlsx"},
            ],
        },
        {
            "shipment_num": 2,
            "order_no": "ORD-INT-002",
            "bl_no": "BL-INT-002",
            "container": "CONT-INT-002",
            "has_planted_discrepancy": True,
            "discrepancy_in_doc_type": "invoice",
            "discrepancy_detail": ["[invoice] order no altered ORD-INT-002 -> ORD-INT-002-ALTERED"],
            "documents": [
                {"file": "S02_BL_v1.docx", "type": "bill_of_lading", "format": "docx"},
                {"file": "S02_Invoice_v1.xlsx", "type": "invoice", "format": "xlsx"},
            ],
        },
    ])

    return corpus, ak_path


# ---------------------------------------------------------------------------
# Integration tests
# ---------------------------------------------------------------------------


class TestIntegrationMixedFormats:
    """End-to-end integration test on a mixed-format corpus."""

    def test_pipeline_runs_without_error(self, mixed_corpus, tmp_path):
        """Full pipeline runs on a mixed docx+xlsx corpus without errors."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        results = run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        assert results["summary"]["shipments"] == 2
        assert results["summary"]["passed"] + results["summary"]["failed"] == 2

    def test_all_stage_artifacts_written(self, mixed_corpus, tmp_path):
        """Every stage's JSON artifact is written to the output directory."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        for artifact in ["raw_docs.json", "canonical_docs.json", "groups.json",
                         "verdicts.json", "results.json"]:
            assert (out_dir / artifact).exists(), f"{artifact} not written"

    def test_correct_doc_count(self, mixed_corpus, tmp_path):
        """All 4 documents are ingested and extracted."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        with open(out_dir / "raw_docs.json") as f:
            raw = json.load(f)
        assert len(raw) == 4

        with open(out_dir / "canonical_docs.json") as f:
            canonical = json.load(f)
        assert len(canonical) == 4

    def test_two_groups_formed(self, mixed_corpus, tmp_path):
        """Documents are correctly grouped into 2 shipment groups."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        with open(out_dir / "groups.json") as f:
            groups = json.load(f)
        assert len(groups) == 2

    def test_clean_shipment_passes(self, mixed_corpus, tmp_path):
        """The clean shipment (S01) gets a PASS verdict."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        results = run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        # Find the group that contains S01 documents
        clean_group = None
        for s in results["shipments"]:
            paths = [d["source_path"] for d in s["documents"]]
            if any("S01" in p for p in paths):
                clean_group = s
                break

        assert clean_group is not None, "S01 group not found"
        assert clean_group["verdict"] == "PASS", \
            f"Clean shipment S01 got FAIL: {clean_group['findings']}"

    def test_broken_shipment_fails(self, mixed_corpus, tmp_path):
        """The broken shipment (S02) gets a FAIL verdict with order_no finding."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        results = run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        # Find the group that contains S02 documents
        broken_group = None
        for s in results["shipments"]:
            paths = [d["source_path"] for d in s["documents"]]
            if any("S02" in p for p in paths):
                broken_group = s
                break

        assert broken_group is not None, "S02 group not found"
        assert broken_group["verdict"] == "FAIL", \
            f"Broken shipment S02 got PASS (should FAIL)"

        # Should have at least one HIGH finding on order_no
        high_findings = [f for f in broken_group["findings"] if f["severity"] == "high"]
        assert len(high_findings) >= 1, \
            f"No HIGH findings for S02: {broken_group['findings']}"
        assert any("order_no" in f["field"] for f in high_findings), \
            f"order_no not in findings: {high_findings}"

    def test_suspect_doc_identified(self, mixed_corpus, tmp_path):
        """The suspect document in the broken shipment is correctly identified."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        results = run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        broken_group = None
        for s in results["shipments"]:
            if s["verdict"] == "FAIL":
                broken_group = s
                break

        assert broken_group is not None
        assert len(broken_group["suspect_doc_ids"]) >= 1, \
            "No suspects identified for broken shipment"

        # The suspect should include the xlsx document (which has the planted error)
        suspect_docs = [
            d for d in broken_group["documents"]
            if d["doc_id"] in broken_group["suspect_doc_ids"]
        ]
        assert len(suspect_docs) >= 1
        # At least one suspect should be the xlsx source (the altered doc)
        suspect_paths = [d.get("source_path", "") for d in suspect_docs]
        assert any(".xlsx" in p for p in suspect_paths), \
            f"xlsx invoice not among suspects: {suspect_paths}"


class TestIntegrationScoring:
    """Integration tests for the scoring harness on the mixed corpus."""

    def test_scoring_matches_answer_key(self, mixed_corpus, tmp_path):
        """Scoring the pipeline results against the answer key produces
        correct metrics."""
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        corpus, ak_path = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        assert scorecard["recall"] == "1/1", \
            f"Recall: {scorecard['recall']}"
        assert scorecard["false_positives"] == 0, \
            f"False positives: {scorecard['false_positives']}"
        assert scorecard["grouping_accuracy"] == "2/2", \
            f"Grouping: {scorecard['grouping_accuracy']}"

    def test_scoring_individual_metrics(self, mixed_corpus, tmp_path):
        """The scoring harness reports correct individual metrics for the mini corpus.

        Note: overall_pass is False because targets are hard-coded for the full
        12-shipment/5-error corpus. We test the individual metrics instead.
        """
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        corpus, ak_path = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        # These individual metrics should be correct for our mini corpus
        assert scorecard["recall"] == "1/1"
        assert scorecard["false_positives"] == 0
        assert scorecard["precision"] == 1.0
        assert scorecard["grouping_accuracy"] == "2/2"

        # overall_pass is False because targets require 5/5 and 12/12
        # (designed for the full corpus, not a 2-shipment mini)

    def test_scoring_confusion_table(self, mixed_corpus, tmp_path):
        """The confusion table has correct entries for both shipments."""
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        corpus, ak_path = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        confusion = scorecard["confusion_table"]
        assert len(confusion) == 2

        # S01: clean, should be PASS
        s01 = next(c for c in confusion if c["shipment"] == 1)
        assert s01["expected_verdict"] == "PASS"
        assert s01["predicted_verdict"] == "PASS"
        assert s01["correct"] is True

        # S02: planted, should be FAIL
        s02 = next(c for c in confusion if c["shipment"] == 2)
        assert s02["expected_verdict"] == "FAIL"
        assert s02["predicted_verdict"] == "FAIL"
        assert s02["correct"] is True

    def test_scorecard_written_to_disk(self, mixed_corpus, tmp_path):
        """The scorecard can be written to disk and reloaded."""
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        corpus, ak_path = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        sc_path = out_dir / "scorecard.json"
        with open(sc_path, "w") as f:
            json.dump(scorecard, f, indent=2)

        assert sc_path.exists()
        loaded = json.loads(sc_path.read_text())
        assert loaded["recall"] == "1/1"
        assert loaded["false_positives"] == 0
        assert loaded["grouping_accuracy"] == "2/2"


class TestIntegrationReportGeneration:
    """Integration tests for the reporting stage."""

    def test_reports_written(self, mixed_corpus, tmp_path):
        """Per-shipment markdown reports are written."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        reports_dir = out_dir / "reports"
        assert reports_dir.exists()
        report_files = list(reports_dir.glob("*.md"))
        assert len(report_files) == 2, f"Expected 2 reports, got {len(report_files)}"

    def test_correction_draft_for_fail(self, mixed_corpus, tmp_path):
        """A correction draft is written for the FAIL shipment."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        corrections_dir = out_dir / "corrections"
        assert corrections_dir.exists()
        correction_files = list(corrections_dir.glob("*.txt"))
        assert len(correction_files) >= 1, "No correction drafts written"

        # The correction should mention the order number
        content = correction_files[0].read_text()
        assert "ORD-INT-002" in content or "order" in content.lower()

    def test_results_json_structure(self, mixed_corpus, tmp_path):
        """results.json has the expected top-level structure."""
        from docverify.pipeline import run_pipeline

        corpus, _ = mixed_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        with open(out_dir / "results.json") as f:
            results = json.load(f)

        assert "summary" in results
        assert "shipments" in results
        assert results["summary"]["shipments"] == 2
        assert results["summary"]["passed"] == 1
        assert results["summary"]["failed"] == 1

        for s in results["shipments"]:
            assert "group_id" in s
            assert "verdict" in s
            assert "identifiers" in s
            assert "documents" in s
            assert "findings" in s


# ---------------------------------------------------------------------------
# Numeric discrepancy corpus fixture
# ---------------------------------------------------------------------------


@pytest.fixture()
def numeric_corpus(tmp_path: Path) -> tuple[Path, Path]:
    """Create a tiny corpus with a planted numeric discrepancy (net_kg).

    Clean shipment (S01):
      - BL (docx): ORD-NUM-001, BL-NUM-001, CONT-NUM-001, net_kg=5000
      - Invoice (xlsx): same values

    Broken shipment (S02):
      - BL (docx): ORD-NUM-002, BL-NUM-002, CONT-NUM-002, net_kg=5000
      - Packing list (docx): same identifiers but net_kg=4900 (planted numeric error)
    """
    corpus = tmp_path / "corpus"
    corpus.mkdir()

    # --- Clean shipment (S01) ---
    _make_docx(
        str(corpus / "S01_BL_v1.docx"),
        "BILL OF LADING",
        {
            "Order No.": "ORD-NUM-001",
            "B/L No.": "BL-NUM-001",
            "Container No.": "CONT-NUM-001",
            "Shipper": "Alpha Foods S.r.l.",
        },
        body="Total net weight: 5000.00 kg\nTotal gross weight: 5500.00 kg\nTotal cartons: 400",
    )

    _make_docx(
        str(corpus / "S01_Invoice_v1.docx"),
        "COMMERCIAL INVOICE",
        {
            "Order No.": "ORD-NUM-001",
            "B/L No.": "BL-NUM-001",
            "Container No.": "CONT-NUM-001",
            "Shipper": "Alpha Foods S.r.l.",
        },
        body="Total net weight: 5000.00 kg\nTotal gross weight: 5500.00 kg\nTotal cartons: 400",
    )

    # --- Broken shipment (S02) — numeric discrepancy in net_kg ---
    _make_docx(
        str(corpus / "S02_BL_v1.docx"),
        "BILL OF LADING",
        {
            "Order No.": "ORD-NUM-002",
            "B/L No.": "BL-NUM-002",
            "Container No.": "CONT-NUM-002",
            "Shipper": "Gamma Pasta Co.",
        },
        body="Total net weight: 5000.00 kg\nTotal gross weight: 5500.00 kg\nTotal cartons: 400",
    )

    _make_docx(
        str(corpus / "S02_PackingList_v1.docx"),
        "PACKING LIST",
        {
            "Order No.": "ORD-NUM-002",
            "B/L No.": "BL-NUM-002",
            "Container No.": "CONT-NUM-002",
            "Shipper": "Gamma Pasta Co.",
        },
        body="Total net weight: 4900.00 kg\nTotal gross weight: 5500.00 kg\nTotal cartons: 400",
    )

    # Answer key
    ak_path = _make_answer_key(tmp_path, [
        {
            "shipment_num": 1,
            "order_no": "ORD-NUM-001",
            "bl_no": "BL-NUM-001",
            "container": "CONT-NUM-001",
            "has_planted_discrepancy": False,
            "documents": [
                {"file": "S01_BL_v1.docx", "type": "bill_of_lading", "format": "docx"},
                {"file": "S01_Invoice_v1.docx", "type": "invoice", "format": "docx"},
            ],
        },
        {
            "shipment_num": 2,
            "order_no": "ORD-NUM-002",
            "bl_no": "BL-NUM-002",
            "container": "CONT-NUM-002",
            "has_planted_discrepancy": True,
            "discrepancy_in_doc_type": "packing_list",
            "discrepancy_detail": ["[packing_list] net_kg altered 5000.00 -> 4900.00"],
            "documents": [
                {"file": "S02_BL_v1.docx", "type": "bill_of_lading", "format": "docx"},
                {"file": "S02_PackingList_v1.docx", "type": "packing_list", "format": "docx"},
            ],
        },
    ])

    return corpus, ak_path


# ---------------------------------------------------------------------------
# Integration tests for numeric discrepancies
# ---------------------------------------------------------------------------


class TestIntegrationNumericDiscrepancy:
    """End-to-end tests verifying the pipeline catches numeric discrepancies."""

    def test_numeric_mismatch_detected(self, numeric_corpus, tmp_path):
        """A shipment with a net_kg mismatch gets a FAIL verdict."""
        from docverify.pipeline import run_pipeline

        corpus, _ = numeric_corpus
        out_dir = tmp_path / "out"
        results = run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        broken_group = None
        for s in results["shipments"]:
            paths = [d["source_path"] for d in s["documents"]]
            if any("S02" in p for p in paths):
                broken_group = s
                break

        assert broken_group is not None, "S02 group not found"
        assert broken_group["verdict"] == "FAIL", \
            f"Numeric mismatch shipment got PASS: {broken_group['findings']}"

    def test_numeric_mismatch_has_high_finding(self, numeric_corpus, tmp_path):
        """The net_kg mismatch produces a HIGH-severity finding."""
        from docverify.pipeline import run_pipeline

        corpus, _ = numeric_corpus
        out_dir = tmp_path / "out"
        results = run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        broken_group = None
        for s in results["shipments"]:
            if s["verdict"] == "FAIL":
                broken_group = s
                break

        assert broken_group is not None
        high_findings = [f for f in broken_group["findings"] if f["severity"] == "high"]
        assert len(high_findings) >= 1, \
            f"No HIGH findings for numeric mismatch: {broken_group['findings']}"
        assert any("net_kg" in f["field"] for f in high_findings), \
            f"net_kg not in HIGH findings: {high_findings}"

    def test_numeric_mismatch_suspect_identified(self, numeric_corpus, tmp_path):
        """The document with the wrong net_kg is identified as suspect."""
        from docverify.pipeline import run_pipeline

        corpus, _ = numeric_corpus
        out_dir = tmp_path / "out"
        results = run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        broken_group = None
        for s in results["shipments"]:
            if s["verdict"] == "FAIL":
                broken_group = s
                break

        assert broken_group is not None
        assert len(broken_group["suspect_doc_ids"]) >= 1, \
            "No suspects identified for numeric mismatch"

        # The suspect should be the packing list (which has the wrong net_kg)
        suspect_docs = [
            d for d in broken_group["documents"]
            if d["doc_id"] in broken_group["suspect_doc_ids"]
        ]
        suspect_paths = [d.get("source_path", "") for d in suspect_docs]
        assert any("PackingList" in p for p in suspect_paths), \
            f"Packing list not among suspects: {suspect_paths}"

    def test_numeric_mismatch_scoring(self, numeric_corpus, tmp_path):
        """Scoring harness correctly evaluates numeric discrepancy localization."""
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        corpus, ak_path = numeric_corpus
        out_dir = tmp_path / "out"
        run_pipeline(str(corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        assert scorecard["recall"] == "1/1", \
            f"Recall: {scorecard['recall']}"
        assert scorecard["false_positives"] == 0, \
            f"False positives: {scorecard['false_positives']}"

        # Check localization detected the correct field
        loc_details = scorecard["localization_details"]
        assert len(loc_details) == 1
        assert loc_details[0]["correct"] is True, \
            f"Localization failed: {loc_details}"
        assert loc_details[0]["expected_field"] == "net_kg", \
            f"Expected field mismatch: {loc_details[0]['expected_field']}"
        assert loc_details[0]["found_field"] == "net_kg", \
            f"Found field mismatch: {loc_details[0]['found_field']}"
