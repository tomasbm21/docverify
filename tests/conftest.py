"""Shared pytest fixtures for docverify tests."""

import pytest


@pytest.fixture
def test_dir(tmp_path):
    """Create a temporary directory with malformed test documents.

    Provides the test_dir fixture expected by test_malformed_docs.py.
    """
    from tests.test_malformed_docs import create_test_files

    # Patch the test_dir path to use tmp_path
    import tests.test_malformed_docs as mod
    orig_path = mod.Path

    class PatchedPath(type(orig_path)):
        """Redirect the test directory to tmp_path."""
        pass

    # Instead of patching Path globally, just create the files directly
    # in the tmp directory by calling the helper with a patched parent
    test_dir = tmp_path / "data" / "test_malformed"
    test_dir.mkdir(parents=True)

    # Create the test files inline (mirrors create_test_files)
    from docx import Document
    from openpyxl import Workbook
    import zipfile
    import io

    # 1. Empty .docx
    empty_docx = Document()
    for p in empty_docx.paragraphs:
        p._element.getparent().remove(p._element)
    empty_docx.save(test_dir / "empty.docx")

    # 2. Header only .docx
    header_only = Document()
    header_only.add_heading("BILL OF LADING", level=1)
    header_only.save(test_dir / "header_only.docx")

    # 3. Empty .xlsx
    wb_empty = Workbook()
    ws_empty = wb_empty.active
    ws_empty.title = "EmptySheet"
    wb_empty.save(test_dir / "empty_sheet.xlsx")

    # 4. Mixed encoding .docx
    mixed_doc = Document()
    mixed_doc.add_paragraph("UTF-8: Hello World")
    mixed_doc.add_paragraph("Latin-1 style: cafe resume")
    mixed_doc.save(test_dir / "mixed_encoding.docx")

    # 5. Corrupted .docx (valid zip, truncated XML)
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml",
                     '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                     '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                     '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                     '<Default Extension="xml" ContentType="application/xml"/>'
                     '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                     "</Types>")
        zf.writestr("_rels/.rels",
                     '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                     '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                     '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
                     "</Relationships>")
        zf.writestr("word/_rels/document.xml.rels",
                     '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                     '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
        zf.writestr("word/document.xml",
                     '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     '<w:body><w:p><w:r><w:t>BILL OF LADING</w:t></w:r></w:p>'
                     "<w:body><w:p><w:r><w:t>Truncated")
    with open(test_dir / "corrupted.docx", "wb") as f:
        f.write(zip_buf.getvalue())

    # 6. Zero-byte .docx
    with open(test_dir / "zero_byte.docx", "wb"):
        pass

    # 7. Fake .docx (plain text)
    with open(test_dir / "fake_docx.docx", "w", encoding="utf-8") as f:
        f.write("This is actually a plain text file pretending to be a docx.\n")
        f.write("BILL OF LADING\nInvoice: 12345\nDate: 2025-01-15\n")

    return test_dir
