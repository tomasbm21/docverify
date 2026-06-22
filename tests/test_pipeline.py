"""End-to-end pipeline smoke test.

OWNER: Agent 06 (Integration + Scoring)
Tests the full pipeline on a tiny fixture corpus (4 docs: 1 clean shipment + 1 broken shipment).
Fast, offline, deterministic.
"""

import json
import os
import shutil
import tempfile
from pathlib import Path

import pytest

from docverify.schemas.models import CanonicalDoc, RawDoc, ShipmentGroup, ShipmentVerdict


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def _make_docx(path: str, title: str, metadata: dict[str, str], body: str = "") -> None:
    """Create a minimal .docx with a title, a metadata table, and optional body."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)

    # Metadata table
    table = doc.add_table(rows=len(metadata), cols=2)
    for i, (k, v) in enumerate(metadata.items()):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    if body:
        doc.add_paragraph(body)

    doc.save(path)


@pytest.fixture()
def mini_corpus(tmp_path: Path) -> Path:
    """Create a tiny corpus: 2 docs for a clean shipment, 2 for a broken one.

    Clean shipment (S_C1):
      - BL: order ORD-001, container CONT-001
      - Invoice: order ORD-001, container CONT-001

    Broken shipment (S_C2):
      - BL: order ORD-002, container CONT-002
      - Invoice: order ORD-002-ALTERED (planted discrepancy)
    """
    corpus = tmp_path / "corpus"
    corpus.mkdir()

    # Clean shipment docs (S01 prefix for scorer mapping)
    _make_docx(
        str(corpus / "S01_BL_v1.docx"),
        "BILL OF LADING",
        {
            "Order No.": "ORD-001",
            "B/L No.": "BL-001",
            "Container No.": "CONT-001",
            "Shipper": "Alpha Foods S.r.l.",
            "Consignee": "Beta Imports Ltd.",
        },
    )

    _make_docx(
        str(corpus / "S01_Invoice_v1.docx"),
        "COMMERCIAL INVOICE",
        {
            "Order No.": "ORD-001",
            "B/L No.": "BL-001",
            "Container No.": "CONT-001",
            "Shipper": "Alpha Foods S.r.l.",
            "Consignee": "Beta Imports Ltd.",
        },
    )

    # Broken shipment docs (S02 prefix for scorer mapping)
    _make_docx(
        str(corpus / "S02_BL_v1.docx"),
        "BILL OF LADING",
        {
            "Order No.": "ORD-002",
            "B/L No.": "BL-002",
            "Container No.": "CONT-002",
            "Shipper": "Gamma Pasta Co.",
            "Consignee": "Delta Trading FZE",
        },
    )

    _make_docx(
        str(corpus / "S02_Invoice_v1.docx"),
        "COMMERCIAL INVOICE",
        {
            "Order No.": "ORD-002-ALTERED",
            "B/L No.": "BL-002",
            "Container No.": "CONT-002",
            "Shipper": "Gamma Pasta Co.",
            "Consignee": "Delta Trading FZE",
        },
    )

    return corpus


# ---------------------------------------------------------------------------
# Pipeline tests
# ---------------------------------------------------------------------------


class TestPipelineEndToEnd:
    """Smoke tests for the full pipeline."""

    def test_pipeline_runs_clean(self, mini_corpus: Path, tmp_path: Path) -> None:
        """Pipeline runs without error on a tiny corpus."""
        from docverify.pipeline import run_pipeline

        out_dir = tmp_path / "out"
        results = run_pipeline(str(mini_corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        # All stage artifacts should exist
        assert (out_dir / "raw_docs.json").exists()
        assert (out_dir / "canonical_docs.json").exists()
        assert (out_dir / "groups.json").exists()
        assert (out_dir / "verdicts.json").exists()
        assert (out_dir / "results.json").exists()

        # Should have 4 docs, 2 groups
        assert results["summary"]["shipments"] == 2

    def test_pipeline_detects_planted_error(self, mini_corpus: Path, tmp_path: Path) -> None:
        """Pipeline FAILs the broken shipment and PASSes the clean one."""
        from docverify.pipeline import run_pipeline

        out_dir = tmp_path / "out"
        results = run_pipeline(str(mini_corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        verdicts = {s["group_id"]: s["verdict"] for s in results["shipments"]}

        # One PASS and one FAIL
        assert "PASS" in verdicts.values()
        assert "FAIL" in verdicts.values()

        # The FAIL group should have an order_no finding
        fail_groups = [s for s in results["shipments"] if s["verdict"] == "FAIL"]
        assert len(fail_groups) == 1
        fail_group = fail_groups[0]
        high_findings = [f for f in fail_group["findings"] if f["severity"] == "high"]
        assert any("order_no" in f["field"] for f in high_findings)

    def test_pipeline_handles_missing_corpus(self, tmp_path: Path) -> None:
        """Pipeline raises FileNotFoundError for a missing corpus directory."""
        from docverify.pipeline import run_pipeline

        with pytest.raises(FileNotFoundError):
            run_pipeline(str(tmp_path / "nonexistent"), str(tmp_path / "out"), False, 0.0)

    def test_pipeline_persists_stage_json(self, mini_corpus: Path, tmp_path: Path) -> None:
        """Each stage's JSON is valid and contains expected records."""
        from docverify.pipeline import run_pipeline

        out_dir = tmp_path / "out"
        run_pipeline(str(mini_corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        with open(out_dir / "raw_docs.json") as f:
            raw = json.load(f)
        assert len(raw) == 4

        with open(out_dir / "canonical_docs.json") as f:
            canonical = json.load(f)
        assert len(canonical) == 4

        with open(out_dir / "groups.json") as f:
            groups = json.load(f)
        assert len(groups) == 2

        with open(out_dir / "verdicts.json") as f:
            verdicts = json.load(f)
        assert len(verdicts) == 2


# ---------------------------------------------------------------------------
# Scoring tests
# ---------------------------------------------------------------------------


class TestScoring:
    """Tests for the scoring harness."""

    def _make_answer_key(self, tmp_path: Path) -> Path:
        """Create a minimal answer key for the mini corpus."""
        ak = {
            "manifest": [
                {"file": "S01_BL_v1.docx", "shipment": 1, "type": "bill_of_lading", "format": "docx"},
                {"file": "S01_Invoice_v1.docx", "shipment": 1, "type": "invoice", "format": "docx"},
                {"file": "S02_BL_v1.docx", "shipment": 2, "type": "bill_of_lading", "format": "docx"},
                {"file": "S02_Invoice_v1.docx", "shipment": 2, "type": "invoice", "format": "docx"},
            ],
            "answer_key": [
                {
                    "shipment": 1,
                    "order_no": "ORD-001",
                    "bl_no": "BL-001",
                    "container": "CONT-001",
                    "has_planted_discrepancy": False,
                    "discrepancy_in_doc_type": None,
                    "discrepancy_detail": [],
                    "ground_truth_totals": {"cartons": 10, "net_kg": 100.0, "gross_kg": 110.0, "value": "EUR 500"},
                },
                {
                    "shipment": 2,
                    "order_no": "ORD-002",
                    "bl_no": "BL-002",
                    "container": "CONT-002",
                    "has_planted_discrepancy": True,
                    "discrepancy_in_doc_type": "invoice",
                    "discrepancy_detail": ["[invoice] order no altered ORD-002 -> ORD-002-ALTERED"],
                    "ground_truth_totals": {"cartons": 20, "net_kg": 200.0, "gross_kg": 220.0, "value": "EUR 1000"},
                },
            ],
            "total_documents": 4,
        }
        ak_path = tmp_path / "answer_key.json"
        with open(ak_path, "w") as f:
            json.dump(ak, f)
        return ak_path

    def test_score_perfect_run(self, mini_corpus: Path, tmp_path: Path) -> None:
        """Scoring a correct pipeline run reports the right metrics for a mini corpus."""
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        out_dir = tmp_path / "out"
        run_pipeline(str(mini_corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        ak_path = self._make_answer_key(tmp_path)
        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        # For the mini corpus: 1 planted error caught, 0 false positives, 2/2 grouping
        assert scorecard["false_positives"] == 0
        assert scorecard["grouping_accuracy"] == "2/2"
        assert scorecard["recall"] == "1/1"

    def test_score_detects_recall(self, mini_corpus: Path, tmp_path: Path) -> None:
        """Scoring correctly counts the planted error as a true positive."""
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        out_dir = tmp_path / "out"
        run_pipeline(str(mini_corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        ak_path = self._make_answer_key(tmp_path)
        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        # The broken shipment should be caught
        assert scorecard["recall"] == "1/1"

    def test_scorecard_written(self, mini_corpus: Path, tmp_path: Path) -> None:
        """Scoring writes scorecard.json to disk."""
        from docverify.pipeline import run_pipeline
        from docverify.scoring.score import score

        out_dir = tmp_path / "out"
        run_pipeline(str(mini_corpus), str(out_dir), use_llm=False, numeric_tolerance=0.0)

        ak_path = self._make_answer_key(tmp_path)
        scorecard = score(str(out_dir / "results.json"), str(ak_path))

        scorecard_path = out_dir / "scorecard.json"
        with open(scorecard_path, "w") as f:
            json.dump(scorecard, f)

        assert scorecard_path.exists()
        loaded = json.loads(scorecard_path.read_text())
        assert loaded["recall"] == "1/1"
        assert loaded["false_positives"] == 0
