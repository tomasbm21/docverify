"""On-the-fly shipping document generator.

Creates realistic .docx and .xlsx shipping documents with randomized data.
Can inject deliberate errors for testing the verification pipeline.

Usage:
    from docverify.agents.generator import generate_shipment_set

    # Generate a complete set of docs for one shipment
    files = generate_shipment_set("/tmp/output", inject_errors=True)
"""

import os
import random
from pathlib import Path

from docverify.utils import get_logger

logger = get_logger(__name__)

# --- Reference data for realistic generation ---

PASTA_PRODUCTS = [
    "Spaghetti n.5 500g", "Penne Rigate 500g", "Fusilli 500g",
    "Linguine 500g", "Farfalle 500g", "Rigatoni 500g",
    "Bucatini 500g", "Orecchiette 500g", "Conchiglie 500g",
    "Spaghetti n.3 5kg", "Penne Rigate 5kg", "Fusilli 5kg",
]

PORTS = [
    ("ITGOA", "Genova, Italy"),
    ("ITNAP", "Napoli, Italy"),
    ("ITSAL", "Salerno, Italy"),
    ("LBBEY", "Beirut, Lebanon"),
    ("TRIST", "Istanbul, Turkey"),
    ("EGALY", "Alexandria, Egypt"),
]

VESSELS = [
    ("MSC ANNA", "V.2024W12"),
    ("CMA CGM MARCO POLO", "V.FAL03E"),
    ("EVER GIVEN", "V.EG045A"),
    ("MAERSK SELETAR", "V.MS089F"),
    ("COSCO SHIPPING ARIES", "V.CS012B"),
]

CURRENCIES = ["EUR", "USD"]


def _random_order_no(rng: random.Random) -> str:
    year = rng.choice([2024, 2025, 2026])
    num = rng.randint(10000, 99999)
    return f"ORD-{year}-{num}"


def _random_bl_no(rng: random.Random) -> str:
    prefix = rng.choice(["MEDU", "CMAU", "MSCU", "EGLV", "COSU"])
    num = rng.randint(1000000, 9999999)
    return f"{prefix}{num}"


def _random_container(rng: random.Random) -> str:
    prefix = rng.choice(["SLNU", "BRLU", "CCLU", "GESU", "TCLU"])
    num = rng.randint(100000, 999999)
    return f"{prefix}{num}"


def _random_seal(rng: random.Random) -> str:
    return str(rng.randint(10000000, 99999999))


def generate_shipment_data(rng_seed: int | None = None) -> dict:
    """Generate realistic shipment data.

    Returns:
        Dict with all fields needed to create BL, packing list, invoice, etc.
    """
    rng = random.Random(rng_seed)

    products = rng.sample(PASTA_PRODUCTS, k=rng.randint(2, 5))
    vessel = rng.choice(VESSELS)
    pol = rng.choice(PORTS[:3])
    pod = rng.choice(PORTS[3:])
    currency = rng.choice(CURRENCIES)

    items = []
    total_cartons = 0
    total_net = 0.0
    total_gross = 0.0
    total_value = 0.0

    for prod in products:
        cartons = rng.randint(50, 500)
        net_per = round(rng.uniform(0.48, 0.52), 3)
        gross_per = round(net_per + rng.uniform(0.02, 0.05), 3)
        net_kg = round(cartons * net_per, 2)
        gross_kg = round(cartons * gross_per, 2)
        unit_price = round(rng.uniform(1.2, 2.5), 2)
        amount = round(cartons * unit_price, 2)

        items.append({
            "description": prod,
            "lot": f"LOT-{rng.randint(1000, 9999)}",
            "cartons": cartons,
            "net_kg": net_kg,
            "gross_kg": gross_kg,
            "unit_price": unit_price,
            "amount": amount,
        })
        total_cartons += cartons
        total_net += net_kg
        total_gross += gross_kg
        total_value += amount

    return {
        "order_no": _random_order_no(rng),
        "bl_no": _random_bl_no(rng),
        "container_no": _random_container(rng),
        "seal_no": _random_seal(rng),
        "reference": f"REF-{rng.randint(1000, 9999)}",
        "shipper": "Garofalo S.p.A. — Pastagarovo Division",
        "shipper_addr": "Via dei Pastai, 15 — 80100 Napoli, Italy",
        "consignee": "Pastagarovo S.a.l.",
        "consignee_addr": "Port Area, Beirut, Lebanon",
        "vessel": vessel[0],
        "voyage": vessel[1],
        "pol": pol[0],
        "pol_name": pol[1],
        "pod": pod[0],
        "pod_name": pod[1],
        "ship_date": f"2026-{rng.randint(1,12):02d}-{rng.randint(1,28):02d}",
        "items": items,
        "totals": {
            "cartons": total_cartons,
            "net_kg": round(total_net, 2),
            "gross_kg": round(total_gross, 2),
            "value": round(total_value, 2),
            "currency": currency,
        },
        "incoterm": rng.choice(["FOB", "CIF", "CFR"]),
    }


def _inject_error(data: dict, rng: random.Random) -> dict:
    """Inject a random error into the shipment data.

    Modifies one field in one document type to create a verifiable discrepancy.
    Returns (modified_data, error_description).
    """
    import copy
    modified = copy.deepcopy(data)

    error_type = rng.choice(["order_no", "container_no", "net_kg", "cartons"])

    if error_type == "order_no":
        # Change last digit of order number
        old = modified["order_no"]
        last_digit = int(old[-1])
        new_digit = (last_digit + rng.randint(1, 5)) % 10
        modified["order_no"] = old[:-1] + str(new_digit)
        return modified, f"order_no: {old} -> {modified['order_no']}"

    elif error_type == "container_no":
        old = modified["container_no"]
        suffix = rng.randint(100000, 999999)
        modified["container_no"] = f"BRLU{suffix}"
        return modified, f"container_no: {old} -> {modified['container_no']}"

    elif error_type == "net_kg":
        # Change net weight by 5-15%
        item_idx = rng.randint(0, len(modified["items"]) - 1)
        old = modified["items"][item_idx]["net_kg"]
        factor = rng.uniform(1.05, 1.15)
        modified["items"][item_idx]["net_kg"] = round(old * factor, 2)
        return modified, f"net_kg[{item_idx}]: {old} -> {modified['items'][item_idx]['net_kg']}"

    elif error_type == "cartons":
        item_idx = rng.randint(0, len(modified["items"]) - 1)
        old = modified["items"][item_idx]["cartons"]
        modified["items"][item_idx]["cartons"] = old + rng.randint(1, 10)
        return modified, f"cartons[{item_idx}]: {old} -> {modified['items'][item_idx]['cartons']}"

    return modified, "no error"


def generate_docx_bill_of_lading(data: dict, output_path: str) -> str:
    """Generate a Bill of Lading .docx file."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    doc.add_heading("BILL OF LADING", level=1)
    doc.add_paragraph(f"B/L No: {data['bl_no']}")
    doc.add_paragraph(f"Order No: {data['order_no']}")
    doc.add_paragraph(f"Reference: {data['reference']}")
    doc.add_paragraph(f"Container No: {data['container_no']}")
    doc.add_paragraph(f"Seal No: {data['seal_no']}")
    doc.add_paragraph("")

    doc.add_paragraph(f"Shipper: {data['shipper']}")
    doc.add_paragraph(f"  {data['shipper_addr']}")
    doc.add_paragraph(f"Consignee: {data['consignee']}")
    doc.add_paragraph(f"  {data['consignee_addr']}")
    doc.add_paragraph("")

    doc.add_paragraph(f"Vessel: {data['vessel']}  Voyage: {data['voyage']}")
    doc.add_paragraph(f"Port of Loading: {data['pol']} — {data['pol_name']}")
    doc.add_paragraph(f"Port of Discharge: {data['pod']} — {data['pod_name']}")
    doc.add_paragraph(f"Date of Shipment: {data['ship_date']}")
    doc.add_paragraph(f"Incoterm: {data['incoterm']}")
    doc.add_paragraph("")

    # Cargo table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Description", "Lot", "Cartons", "Net KG", "Gross KG"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for item in data["items"]:
        row = table.add_row()
        row.cells[0].text = item["description"]
        row.cells[1].text = item["lot"]
        row.cells[2].text = str(item["cartons"])
        row.cells[3].text = str(item["net_kg"])
        row.cells[4].text = str(item["gross_kg"])

    doc.add_paragraph("")
    doc.add_paragraph(f"Total Cartons: {data['totals']['cartons']}")
    doc.add_paragraph(f"Total Net Weight: {data['totals']['net_kg']} KG")
    doc.add_paragraph(f"Total Gross Weight: {data['totals']['gross_kg']} KG")

    filepath = Path(output_path) / f"BL_{data['order_no']}.docx"
    doc.save(str(filepath))
    return str(filepath)


def generate_docx_packing_list(data: dict, output_path: str) -> str:
    """Generate a Packing List .docx file."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    doc.add_heading("PACKING LIST", level=1)
    doc.add_paragraph(f"Order No: {data['order_no']}")
    doc.add_paragraph(f"Reference: {data['reference']}")
    doc.add_paragraph(f"Container No: {data['container_no']}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Description", "Lot", "Cartons", "Net KG", "Gross KG", "Unit Price"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for item in data["items"]:
        row = table.add_row()
        row.cells[0].text = item["description"]
        row.cells[1].text = item["lot"]
        row.cells[2].text = str(item["cartons"])
        row.cells[3].text = str(item["net_kg"])
        row.cells[4].text = str(item["gross_kg"])
        row.cells[5].text = f"{item['unit_price']:.2f}"

    doc.add_paragraph("")
    doc.add_paragraph(f"Total Cartons: {data['totals']['cartons']}")
    doc.add_paragraph(f"Total Net Weight: {data['totals']['net_kg']} KG")
    doc.add_paragraph(f"Total Gross Weight: {data['totals']['gross_kg']} KG")

    filepath = Path(output_path) / f"PackingList_{data['order_no']}.docx"
    doc.save(str(filepath))
    return str(filepath)


def generate_xlsx_invoice(data: dict, output_path: str, label: str = "Invoice") -> str:
    """Generate an Invoice .xlsx file."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = label

    # Header
    ws["A1"] = f"COMMERCIAL {label.upper()}"
    ws["A1"].font = openpyxl.styles.Font(bold=True, size=14)

    ws["A3"] = "Invoice No:"
    ws["B3"] = f"INV-{data['order_no'].split('-')[-1]}"
    ws["A4"] = "Order No:"
    ws["B4"] = data["order_no"]
    ws["A5"] = "Reference:"
    ws["B5"] = data["reference"]
    ws["A6"] = "Container No:"
    ws["B6"] = data["container_no"]
    ws["A7"] = "Shipper:"
    ws["B7"] = data["shipper"]
    ws["A8"] = "Consignee:"
    ws["B8"] = data["consignee"]

    # Items table
    row = 10
    headers = ["Description", "Lot", "Cartons", "Net KG", "Gross KG", "Unit Price", "Amount"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=h)
        cell.font = openpyxl.styles.Font(bold=True)

    for item in data["items"]:
        row += 1
        ws.cell(row=row, column=1, value=item["description"])
        ws.cell(row=row, column=2, value=item["lot"])
        ws.cell(row=row, column=3, value=item["cartons"])
        ws.cell(row=row, column=4, value=item["net_kg"])
        ws.cell(row=row, column=5, value=item["gross_kg"])
        ws.cell(row=row, column=6, value=item["unit_price"])
        ws.cell(row=row, column=7, value=item["amount"])

    # Totals
    row += 2
    ws.cell(row=row, column=1, value="TOTALS").font = openpyxl.styles.Font(bold=True)
    ws.cell(row=row, column=3, value=data["totals"]["cartons"])
    ws.cell(row=row, column=4, value=data["totals"]["net_kg"])
    ws.cell(row=row, column=5, value=data["totals"]["gross_kg"])
    ws.cell(row=row, column=7, value=data["totals"]["value"])
    ws.cell(row=row + 1, column=7, value=data["totals"]["currency"])

    filepath = Path(output_path) / f"{label}_{data['order_no']}.xlsx"
    wb.save(str(filepath))
    return str(filepath)


def generate_shipment_set(
    output_dir: str,
    rng_seed: int | None = None,
    inject_errors: bool = False,
) -> dict:
    """Generate a complete set of shipping documents for one shipment.

    Args:
        output_dir: Directory to write the generated files.
        rng_seed: Seed for reproducible generation.
        inject_errors: If True, inject a deliberate error into one document.

    Returns:
        Dict with 'files' (list of paths), 'data' (shipment data),
        'error' (error description if injected, else None).
    """
    os.makedirs(output_dir, exist_ok=True)
    rng = random.Random(rng_seed)

    data = generate_shipment_data(rng_seed)

    error_desc = None
    errored_data = data
    if inject_errors:
        errored_data, error_desc = _inject_error(data, rng)

    files = []
    files.append(generate_docx_bill_of_lading(data, output_dir))
    files.append(generate_docx_packing_list(data, output_dir))
    # If errors injected, the invoice uses the errored data so there is a mismatch
    files.append(generate_xlsx_invoice(errored_data, output_dir))

    return {
        "files": files,
        "data": data,
        "error": error_desc,
    }


if __name__ == "__main__":
    import sys
    out = sys.argv[1] if len(sys.argv) > 1 else "data/generated"
    result = generate_shipment_set(out, rng_seed=42, inject_errors=False)
    print(f"Generated {len(result['files'])} files in {out}/")
    for f in result["files"]:
        print(f"  {f}")
