"""Ingestion / Format Normalizer -- reads .docx/.xlsx/.pdf files into RawDoc records.

OWNER: Agent 01 (Ingestion)
FROZEN signatures per CONTRACTS.md §5.

Reads each file, extracts plain text and structured tables, builds a content-hash
doc_id, and returns a RawDoc. Lossless enough that Agent B can extract every field.
"""

from __future__ import annotations

import os
from pathlib import Path

from docverify.schemas.models import RawDoc, SourceFormat
from docverify.utils import content_hash, get_logger

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# .docx reader
# ---------------------------------------------------------------------------

def _read_docx(path: str) -> tuple[str, list[list[list[str]]]]:
    """Extract text and tables from a .docx file.

    Returns (text, tables) where tables is a list of tables, each table a list
    of rows, each row a list of cell strings.  Paragraph text is emitted in
    document order with original spacing preserved (critical for legacy
    fixed-width layouts).
    """
    from docx import Document

    doc = Document(path)

    # Build a mapping of (table element) -> table data so we can interleave
    # paragraphs and tables in document order.
    table_map: dict[int, list[list[str]]] = {}
    for idx, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        table_map[id(table._element)] = rows

    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    # Walk the document body XML in order to interleave paragraphs and tables.
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            # Paragraph
            para_text = child.text or ""
            # python-docx paragraphs don't always have .text on the element;
            # fall back to iterating runs.
            if not para_text:
                # Reconstruct from the Document's paragraph objects
                for para in doc.paragraphs:
                    if para._element is child:
                        para_text = para.text
                        break
            text_parts.append(para_text)
        elif tag == "tbl":
            tbl_data = table_map.get(id(child))
            if tbl_data is not None:
                tables.append(tbl_data)
                # Also add a text representation of the table so the text field
                # contains all content for hashing and downstream extraction.
                for row in tbl_data:
                    text_parts.append("\t".join(row))

    text = "\n".join(text_parts)
    return text, tables


# ---------------------------------------------------------------------------
# .xlsx reader
# ---------------------------------------------------------------------------

def _read_xlsx(path: str) -> tuple[str, list[list[list[str]]]]:
    """Extract text and tables from an .xlsx file.

    Reads each sheet with data_only=True so formula cells return their cached
    value.  Handles merged cells by propagating the top-left value to the
    entire range.  Returns (text, tables) where tables is a list of tables
    (one per sheet).
    """
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)

    all_tables: list[list[list[str]]] = []
    text_parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Build a map of merged cell ranges so we can fill in blanks.
        merge_map: dict[tuple[int, int], str] = {}
        for merge_range in ws.merged_cells.ranges:
            min_row, min_col = merge_range.min_row, merge_range.min_col
            top_left_val = ws.cell(min_row, min_col).value
            top_left_str = "" if top_left_val is None else str(top_left_val)
            for row_idx in range(merge_range.min_row, merge_range.max_row + 1):
                for col_idx in range(merge_range.min_col, merge_range.max_col + 1):
                    merge_map[(row_idx, col_idx)] = top_left_str

        sheet_rows: list[list[str]] = []
        for row_idx in range(1, ws.max_row + 1):
            row_cells: list[str] = []
            for col_idx in range(1, ws.max_column + 1):
                val = ws.cell(row_idx, col_idx).value
                if val is None and (row_idx, col_idx) in merge_map:
                    val = merge_map[(row_idx, col_idx)]
                row_cells.append("" if val is None else str(val))
            # Skip rows that are entirely empty
            if any(c for c in row_cells):
                sheet_rows.append(row_cells)

        if sheet_rows:
            all_tables.append(sheet_rows)

            # Build text representation
            text_parts.append(f"[Sheet: {sheet_name}]")
            for row in sheet_rows:
                # Detect metadata-style rows (label in col A, value in col B,
                # rest empty) and format as "label: value".
                non_empty = [(i, c) for i, c in enumerate(row) if c]
                if len(non_empty) == 2 and non_empty[0][0] == 0 and non_empty[1][0] == 1:
                    text_parts.append(f"{non_empty[0][1]}: {non_empty[1][1]}")
                else:
                    text_parts.append("\t".join(row))

    text = "\n".join(text_parts)
    return text, all_tables


# ---------------------------------------------------------------------------
# .pdf reader
# ---------------------------------------------------------------------------

def _read_pdf(path: str) -> tuple[str, list[list[list[str]]]]:
    """Extract text and tables from a .pdf file.

    Uses pdfplumber for both text and table extraction.  Handles multi-page
    documents by concatenating text across pages and collecting tables from
    every page.  PDFs with no extractable text layer (e.g. scanned images)
    return empty text with a warning — OCR is a separate concern.

    Returns (text, tables) where tables is a list of tables, each table a list
    of rows, each row a list of cell strings.
    """
    import pdfplumber

    text_parts: list[str] = []
    all_tables: list[list[list[str]]] = []
    has_text = False

    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # --- text ---
            page_text = page.extract_text()
            if page_text:
                has_text = True
                text_parts.append(page_text)

            # --- tables ---
            page_tables = page.extract_tables()
            if page_tables:
                for tbl in page_tables:
                    # pdfplumber returns list[list[str | None]]; normalize None -> ""
                    cleaned: list[list[str]] = []
                    for row in tbl:
                        cleaned.append([cell if cell is not None else "" for cell in row])
                    # Skip tables that are entirely empty
                    if any(any(c for c in row) for row in cleaned):
                        all_tables.append(cleaned)
                        # Mirror _read_docx: add tabular text to text_parts
                        for row in cleaned:
                            text_parts.append("\t".join(row))

    if not has_text:
        logger.warning("PDF has no extractable text layer (possibly scanned): %s", path)

    text = "\n".join(text_parts)
    return text, all_tables


# ---------------------------------------------------------------------------
# Public API (FROZEN signatures per CONTRACTS.md §5)
# ---------------------------------------------------------------------------

def ingest_file(path: str) -> RawDoc:
    """Read a single .docx, .xlsx, or .pdf file and return a RawDoc record."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        text, tables = _read_docx(path)
        source_format = SourceFormat.docx
    elif ext == ".xlsx":
        text, tables = _read_xlsx(path)
        source_format = SourceFormat.xlsx
    elif ext == ".pdf":
        text, tables = _read_pdf(path)
        source_format = SourceFormat.pdf
    else:
        raise ValueError(f"Unsupported file extension: {ext!r} (expected .docx, .xlsx, or .pdf)")

    doc_id = content_hash(text, source_path=path)
    return RawDoc(
        doc_id=doc_id,
        source_path=path,
        source_format=source_format,
        text=text,
        tables=tables,
    )


def ingest_dir(corpus_dir: str) -> list[RawDoc]:
    """Read all .docx/.xlsx/.pdf files from corpus_dir and return a list of RawDoc records.

    Skips non-document files quietly.  Never crashes on one bad file — logs a
    warning and continues.
    """
    results: list[RawDoc] = []
    corpus_path = Path(corpus_dir)
    if not corpus_path.is_dir():
        logger.warning("corpus_dir does not exist: %s", corpus_dir)
        return results

    for fpath in sorted(corpus_path.iterdir()):
        if fpath.is_dir():
            continue
        ext = fpath.suffix.lower()
        if ext not in (".docx", ".xlsx", ".pdf"):
            continue
        try:
            raw = ingest_file(str(fpath))
            results.append(raw)
        except Exception as e:
            logger.warning("Skipping %s: not a valid document (%s)", fpath.name, e)
            logger.debug("Full traceback for %s:", fpath.name, exc_info=True)

    return results


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import argparse
    import json
    import sys

    parser = argparse.ArgumentParser(description="Ingest .docx/.xlsx/.pdf files into RawDoc JSON")
    parser.add_argument("corpus_dir", help="Path to the corpus directory")
    parser.add_argument(
        "-o", "--output", default="data/out/raw_docs.json",
        help="Output JSON path (default: data/out/raw_docs.json)",
    )
    args = parser.parse_args()

    # BUG-016: validate corpus path before ingesting
    corpus_path = Path(args.corpus_dir)
    if not corpus_path.exists():
        print(f"Error: path does not exist: {args.corpus_dir}", file=sys.stderr)
        sys.exit(1)
    if not corpus_path.is_dir():
        print(f"Error: path is not a directory: {args.corpus_dir}", file=sys.stderr)
        sys.exit(1)

    docs = ingest_dir(args.corpus_dir)

    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump([d.model_dump(mode="json") for d in docs], f, ensure_ascii=False, indent=2)

    print(f"Wrote {len(docs)} RawDoc records to {out_path}")
